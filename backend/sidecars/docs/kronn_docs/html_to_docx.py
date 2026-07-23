"""Render HTML faithfully into a visually fixed Word document."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pypdfium2 as pdfium
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from weasyprint import CSS, HTML


_PDF_POINTS_PER_INCH = 72
_RENDER_DPI = 150
_EMU_PER_POINT = 12_700


def _configure_section(section, width_points: float, height_points: float) -> None:
    section.page_width = round(width_points * _EMU_PER_POINT)
    section.page_height = round(height_points * _EMU_PER_POINT)
    section.top_margin = 0
    section.right_margin = 0
    section.bottom_margin = 0
    section.left_margin = 0
    section.header_distance = 0
    section.footer_distance = 0


def _configure_image_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)


def render_html_to_docx(
    html: str,
    output_path: str | Path,
    *,
    page_stylesheet: str = "@page { size: A4; margin: 0; }",
) -> None:
    """Render each PDF page as a full-page Word image.

    Word's layout model cannot reproduce arbitrary HTML/CSS. Using the same
    WeasyPrint render as PDF keeps gradients, variables, flex layouts,
    positioned elements, and page breaks visually consistent.
    """
    pdf_bytes = HTML(string=html).write_pdf(
        stylesheets=[CSS(string=page_stylesheet)],
    )
    document = Document()

    with pdfium.PdfDocument(pdf_bytes) as pdf:
        for page_index in range(len(pdf)):
            page = pdf[page_index]
            try:
                width_points, height_points = page.get_size()
                if page_index == 0:
                    _configure_section(
                        document.sections[0],
                        width_points,
                        height_points,
                    )
                    paragraph = document.add_paragraph()
                else:
                    paragraph = document.add_paragraph()
                    paragraph.paragraph_format.page_break_before = True
                _configure_image_paragraph(paragraph)

                bitmap = page.render(
                    scale=_RENDER_DPI / _PDF_POINTS_PER_INCH,
                    fill_color=(255, 255, 255, 255),
                    optimize_mode="print",
                )
                try:
                    image = bitmap.to_pil()
                    image_bytes = BytesIO()
                    image.save(image_bytes, format="PNG", optimize=True)
                    image_bytes.seek(0)
                    paragraph.add_run().add_picture(
                        image_bytes,
                        width=round(width_points * _EMU_PER_POINT),
                    )
                finally:
                    bitmap.close()
            finally:
                page.close()

    document.save(str(output_path))
