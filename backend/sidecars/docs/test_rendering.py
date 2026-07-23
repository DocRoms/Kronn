from __future__ import annotations

from io import BytesIO
import tempfile
import unittest
import zipfile
from pathlib import Path

import pypdfium2 as pdfium
from docx import Document
from PIL import Image, ImageChops
from weasyprint import CSS, HTML

from kronn_docs.html_to_docx import render_html_to_docx
from kronn_docs.server import _pdf_page_stylesheet


STYLED_HTML = """<!doctype html>
<html>
<head>
<style>
  :root { --bg: #0b1020; --card: #121a2f; --text: #eef3ff; }
  html, body { margin: 0; min-height: 100%; background: var(--bg); color: var(--text); }
  .wrap { padding: 28px 20px; }
  .hero {
    background: linear-gradient(135deg, #6d5dfc, #38bdf8);
    border: 2px solid #f472b6;
    border-radius: 18px;
    padding: 22px;
  }
  .cards { display: flex; gap: 12px; margin-top: 20px; }
  .card { flex: 1; min-height: 120px; background: var(--card); padding: 12px; }
</style>
</head>
<body>
  <main class="wrap">
    <header class="hero"><h1>Styled report</h1></header>
    <section class="cards"><div class="card">One</div><div class="card">Two</div></section>
  </main>
</body>
</html>"""


class PdfPageStylesheetTests(unittest.TestCase):
    def test_default_page_has_no_margin(self) -> None:
        css = _pdf_page_stylesheet("<html><body>x</body></html>", "A4")
        self.assertIn("size: A4", css)
        self.assertIn("margin: 0", css)

    def test_author_page_margin_is_preserved(self) -> None:
        html = "<style>@page { margin: 7mm 9mm; }</style>"
        css = _pdf_page_stylesheet(html, "A4 landscape")
        self.assertIn("size: A4 landscape", css)
        self.assertNotIn("margin: 0", css)


class StyledDocxTests(unittest.TestCase):
    def test_embedded_word_page_matches_the_pdf_render(self) -> None:
        stylesheet = "@page { size: A4; margin: 0; }"
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "styled.docx"
            render_html_to_docx(
                STYLED_HTML,
                output,
                page_stylesheet=stylesheet,
            )
            document = Document(output)
            with zipfile.ZipFile(output) as archive:
                image_names = [
                    name
                    for name in archive.namelist()
                    if name.startswith("word/media/")
                ]
                self.assertEqual(len(image_names), 1)
                embedded_image = Image.open(
                    BytesIO(archive.read(image_names[0]))
                ).convert("RGB")

        reference_pdf = HTML(string=STYLED_HTML).write_pdf(
            stylesheets=[CSS(string=stylesheet)]
        )
        with pdfium.PdfDocument(reference_pdf) as pdf:
            page = pdf[0]
            try:
                bitmap = page.render(
                    scale=150 / 72,
                    fill_color=(255, 255, 255, 255),
                    optimize_mode="print",
                )
                try:
                    reference_image = bitmap.to_pil().convert("RGB")
                finally:
                    bitmap.close()
            finally:
                page.close()

        self.assertEqual(document.sections[0].top_margin, 0)
        self.assertEqual(document.sections[0].right_margin, 0)
        self.assertEqual(document.sections[0].bottom_margin, 0)
        self.assertEqual(document.sections[0].left_margin, 0)
        self.assertEqual(len(document.inline_shapes), 1)
        self.assertEqual(embedded_image.size, reference_image.size)
        self.assertIsNone(
            ImageChops.difference(embedded_image, reference_image).getbbox()
        )

    def test_each_rendered_page_gets_its_own_word_page(self) -> None:
        html = """
        <style>
          html, body { margin: 0; background: #123456; }
          .next { break-before: page; }
        </style>
        <div>First page</div><div class="next">Second page</div>
        """
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "multipage.docx"
            render_html_to_docx(html, output)
            document = Document(output)

        self.assertEqual(len(document.inline_shapes), 2)
        self.assertEqual(len(document.paragraphs), 2)
        self.assertTrue(document.paragraphs[1].paragraph_format.page_break_before)


if __name__ == "__main__":
    unittest.main()
